# -*- coding: cp936 -*-
__author__ = 'corey'
# -*- coding:utf-8 -*-
import docx

def printdoc():
    
    document = docx.Document("./1.docx")
    #打印每个段落的内容
    for paragraph in document.paragraphs:
        print paragraph.text
    #打印每个表格的内容
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                print cell.text
        
printdoc()
